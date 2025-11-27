import os
import re
import datetime
from fpdf import FPDF
from docx import Document

# ==========================================================
# 🟦 RÉCUPÉRATION AUTOMATIQUE DU CHEMIN DU DOSSIER /IA/fonts
# ==========================================================
BASE_DIR = os.path.dirname(os.path.abspath(__file__))          # backend/IA
FONT_DIR = os.path.join(BASE_DIR, "fonts")                     # backend/IA/fonts
EXPORT_DIR = os.path.join(BASE_DIR, "exports")                 # backend/IA/exports

os.makedirs(EXPORT_DIR, exist_ok=True)

FONT_REGULAR = os.path.join(FONT_DIR, "DejaVuSans.ttf")
FONT_BOLD = os.path.join(FONT_DIR, "DejaVuSans-Bold.ttf")
FONT_ITALIC = os.path.join(FONT_DIR, "DejaVuSans-Oblique.ttf")

# Vérifier que la police existe (sinon erreur claire)
if not os.path.exists(FONT_REGULAR):
    raise FileNotFoundError(f"Police introuvable : {FONT_REGULAR}")


# ==========================================================
# 🟨 CLASS PDF PERSONNALISÉE
# ==========================================================
class PDF(FPDF):
    def __init__(self):
        super().__init__()
        # Ajout des polices DejaVu (UTF-8)
        self.add_font("DejaVu", "", FONT_REGULAR, uni=True)
        self.add_font("DejaVu", "B", FONT_BOLD if os.path.exists(FONT_BOLD) else FONT_REGULAR, uni=True)
        self.add_font("DejaVu", "I", FONT_ITALIC if os.path.exists(FONT_ITALIC) else FONT_REGULAR, uni=True)

    def header(self):
        BLUE = (30, 55, 153)

        # Ligne horizontale
        self.set_draw_color(*BLUE)
        self.set_line_width(0.8)
        self.line(10, 10, 200, 10)
        self.ln(4)

        # Titre
        self.set_text_color(*BLUE)
        self.set_font("DejaVu", "B", 15)
        self.cell(0, 8, "MEETRECAP - Compte rendu automatique", ln=True)

        # Date
        self.set_text_color(80, 80, 80)
        self.set_font("DejaVu", "", 11)
        today = datetime.datetime.now().strftime("%d/%m/%Y")
        self.cell(
            0, 6,
            f"Date : {today}          Produit par MeetRecap AI",
            ln=True
        )

        self.ln(3)

        # Ligne de séparation
        self.set_draw_color(*BLUE)
        self.set_line_width(0.8)
        self.line(10, 28, 200, 28)
        self.ln(8)


# ==========================================================
# 🟩 UTILITAIRES
# ==========================================================
def is_separator_line(line: str) -> bool:
    return bool(re.match(r'^[\s\u2500-\u257F\-—_]+$', line.strip()))


def clean_lines(text: str) -> list:
    return [
        line for line in text.split("\n")
        if line.strip() and not is_separator_line(line)
    ]


# ==========================================================
# 🟥 GENERATION PDF
# ==========================================================
def save_as_pdf(text: str, output_filename: str):
    pdf_path = os.path.join(EXPORT_DIR, output_filename)

    # Supprimer ancien fichier si présent
    if os.path.exists(pdf_path):
        os.remove(pdf_path)

    pdf = PDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    # Extraire les sections "====="
    raw_sections = text.split("=====")
    sections = []

    for sec in raw_sections:
        sec = sec.strip()
        if not sec:
            continue

        lines = sec.split("\n")
        name = lines[0].strip()
        body = "\n".join(lines[1:]).strip()

        sections.append((name, body))

    BLUE = (30, 55, 153)

    for section_name, section_body in sections:
        # Titre
        pdf.set_text_color(*BLUE)
        pdf.set_font("DejaVu", "B", 13)
        pdf.cell(0, 8, section_name, ln=True)

        # Ligne
        pdf.set_text_color(60, 60, 60)
        pdf.set_font("DejaVu", "", 12)
        pdf.cell(0, 5, "────────────────────", ln=True)
        pdf.ln(4)

        # Contenu
        pdf.set_text_color(0, 0, 0)
        pdf.set_font("DejaVu", "", 12)

        cleaned_lines = clean_lines(section_body)

        for line in cleaned_lines:
            pdf.multi_cell(0, 7, line)
            pdf.ln(1)

        pdf.ln(5)

    pdf.output(pdf_path)
    print(f"📄 PDF généré : {pdf_path}")

    return pdf_path


# ==========================================================
# 🟦 GENERATION WORD
# ==========================================================
def save_as_word(text: str, output_filename: str):
    doc_path = os.path.join(EXPORT_DIR, output_filename)

    doc = Document()
    doc.add_heading("MEETRECAP - Compte-rendu automatique", level=1)

    for line in text.split("\n"):
        if not is_separator_line(line):
            doc.add_paragraph(line)

    doc.save(doc_path)
    print(f"📄 DOCX généré : {doc_path}")

    return doc_path


# ==========================================================
# 🟪 FONCTION PRINCIPALE APPELÉE PAR pipeline_service.py
# ==========================================================
def save_files(text: str, filename: str = "compte_rendu_reunion"):
    pdf_name = f"{filename}.pdf"
    docx_name = f"{filename}.docx"

    pdf_path = save_as_pdf(text, pdf_name)
    docx_path = save_as_word(text, docx_name)

    return {
        "pdf": pdf_path,
        "docx": docx_path
    }
